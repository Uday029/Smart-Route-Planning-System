import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add empty lines for vertical centering
for _ in range(5):
    doc.add_paragraph()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SMART ROUTE PLANNING SYSTEM")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A Desktop Geographical Routing and Emergency Search Engine")
run.italic = True
run.font.size = Pt(16)

doc.add_paragraph('\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A Project Report\nSubmitted in partial fulfillment of the requirements for\nPlacement Preparation Classes")
run.font.size = Pt(14)

doc.add_paragraph('\n\n\n')

# Create a 2-column, 1-row table for submitted by/to
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Cell 1 (Left)
cell1 = table.cell(0, 0)
p1 = cell1.paragraphs[0]
p1.add_run("Submitted By:\n").bold = True
p1.add_run("Udayveer Singh Chaudhary\n")
p1.add_run("Course: B.Tech CSE\n")
p1.add_run("Registration Number: 12320106")

# Cell 2 (Right)
cell2 = table.cell(0, 1)
p2 = cell2.paragraphs[0]
p2.add_run("Submitted To:\n").bold = True
p2.add_run("Sachin Garg\n")
p2.add_run("Placement Preparation Instructor")

# Force Page Break
doc.add_page_break()

# Helper for headings
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_paragraph(text):
    p = doc.add_paragraph()
    p.add_run(text)

# Content
add_heading("1. Abstract")
add_paragraph("The Smart Route Planning System is an advanced desktop application designed to solve complex navigation, logistics, and geographical routing problems. Built entirely in Java using the Swing framework and the modern FlatLaf UI engine, the system models raw geographical data (cities, highways, and expressive coordinates) as an undirected weighted graph.\n")
add_paragraph("Instead of relying on commercial APIs, this project implements core computer science algorithms from scratch. By leveraging advanced Data Structures and Algorithms (DSA)—including Dijkstra’s Algorithm, A* Search, Breadth-First Search (BFS), Depth-First Search (DFS), and Floyd-Warshall—the software calculates optimal travel paths based on dynamic criteria.\n")
add_paragraph("Beyond static pathfinding, the application simulates real-world constraints such as live traffic congestion and weather-induced road closures by dynamically recalculating edge weights on the fly. Furthermore, it incorporates a localized spatial search engine utilizing Priority Queues (Max-Heaps) and Haversine distance calculations to locate nearby emergency services like hospitals, police stations, and petrol pumps within a strict 50km radius. This project serves as a comprehensive demonstration of applying theoretical Computer Science concepts to practical, real-world software engineering challenges.")

doc.add_page_break()

add_heading("2. Introduction")
add_heading("2.1 Background", level=2)
add_paragraph("In the modern era of logistics, emergency response, and daily travel, calculating the most efficient path between two locations is a critical requirement. While commercial solutions like Google Maps and Waze dominate the consumer market, understanding and implementing the underlying routing algorithms from scratch is a fundamental challenge in computer science. Algorithms that solve the Shortest Path Problem are heavily utilized in network routing, supply chain optimization, and artificial intelligence.")

add_heading("2.2 Problem Statement", level=2)
add_paragraph("Traditional static maps fail to account for dynamic variables such as traffic, weather, and localized emergency needs. A robust system must not only find the shortest path but also adapt to changing edge weights (e.g., increased travel time due to a storm) and provide spatial awareness of the immediate surroundings without relying on constant internet access to paid APIs.")

add_heading("2.3 Objectives", level=2)
doc.add_paragraph("• To design and implement a robust Graph data structure capable of representing cities (vertices) and roads (edges) with associated weights.", style='List Bullet')
doc.add_paragraph("• To implement and compare multiple pathfinding algorithms to understand algorithmic efficiency.", style='List Bullet')
doc.add_paragraph("• To develop a spatial search algorithm capable of finding the K-nearest points of interest using heuristic distance formulas.", style='List Bullet')
doc.add_paragraph("• To wrap the mathematical logic in a user-friendly, responsive, and modern Graphical User Interface (GUI).", style='List Bullet')

doc.add_page_break()

add_heading("3. Feasibility Study")
add_heading("3.1 Technical Feasibility", level=2)
add_paragraph("The project strictly requires Java Development Kit (JDK) 17. Since Java is platform-independent, the application can run natively on Windows, macOS, and Linux. The use of Java Swing (a built-in library) mixed with FlatLaf (a lightweight 1MB dependency) ensures that the hardware requirements are extremely low. The project is technically highly feasible.")

add_heading("3.2 Economic Feasibility", level=2)
add_paragraph("This software was developed using entirely open-source tools (Visual Studio Code, OpenJDK, Git, Maven). There are no licensing costs, API usage limits, or server hosting fees required to run the local desktop application. Therefore, the project is 100% economically feasible.")

add_heading("3.3 Operational Feasibility", level=2)
add_paragraph("The end-user is not expected to understand Graph Theory to use the application. The GUI abstractions—such as dropdown menus for cities and simple toggle switches for 'Live Traffic'—ensure that anyone who can use a standard PC can operate the Route Planner.")

doc.add_page_break()

add_heading("4. Technology Stack and Architecture")
add_heading("4.1 Software Technologies", level=2)
doc.add_paragraph("Programming Language: Java (JDK 17) - Chosen for its strong object-oriented capabilities, memory management, and robust standard library for collections.", style='List Bullet')
doc.add_paragraph("GUI Framework: Java Swing - Utilized for constructing the desktop window, layout management, and event handling.", style='List Bullet')
doc.add_paragraph("UI Rendering Engine: FlatLaf - A modern, open-source library that replaces the dated default Java styling with a sleek, minimalist, high-DPI aware interface.", style='List Bullet')
doc.add_paragraph("Architecture Pattern: Model-View-Controller (MVC) conceptual design to separate data structures from visual rendering.", style='List Bullet')

add_heading("4.2 Code Structure and Working Logic", level=2)
add_paragraph("A. Model Layer (City, Graph, Road): The City object acts as a Vertex. The Graph object initializes a HashMap mapping every City ID to a List of Edge objects. When a new road is added, the Graph updates the list.")
add_paragraph("B. View Layer (MainFrame, MapPanel): The MapPanel extends JPanel and overrides paintComponent. It mathematical normalizes geographical Latitude/Longitude pairs into (X, Y) screen pixels, and draws lines (roads) between them.")
add_paragraph("C. Controller Layer (RoutingAlgorithms): These classes contain static methods. They accept the Graph data structure, process the math using Heaps, and return a RouteResult back to the UI.")

doc.add_page_break()

add_heading("5. In-Depth Algorithms and Data Structures")
add_heading("5.1 Core Data Structures Used", level=2)
doc.add_paragraph("Adjacency List: Used to store the map. It ensures that memory scales linearly O(V + E), making it highly efficient for sparse graphs like road networks.", style='List Bullet')
doc.add_paragraph("Priority Queue (Min-Heap): Used in Dijkstra and A* to continuously extract the node with the lowest provisional distance in O(log V) time.", style='List Bullet')
doc.add_paragraph("Priority Queue (Max-Heap): Used in the Nearby Places engine to find the K-nearest hospitals. Maintaining a bounded Max-Heap takes O(N log K) time.", style='List Bullet')

add_heading("5.2 Working Logic of Algorithms", level=2)
doc.add_paragraph("Dijkstra's Algorithm: Initializes all distances to infinity. The source is set to 0. It uses a Min-Heap to evaluate the closest unvisited node. For every neighbor, if the current known distance + the edge weight is less than the neighbor's stored distance, the neighbor's distance is updated and pushed back into the Min-Heap.", style='List Bullet')
doc.add_paragraph("A* (A-Star) Algorithm: Extends Dijkstra by introducing a heuristic h(n). Instead of blindly expanding the closest node, it minimizes f(n) = g(n) + h(n), where g(n) is the cost so far, and h(n) is the estimated Haversine distance to the destination. ", style='List Bullet')
doc.add_paragraph("Breadth-First Search (BFS): Utilizes a standard FIFO Queue (First-In, First-Out). It explores the graph radially layer-by-layer. Edge weights are completely ignored to find paths with the minimum stops.", style='List Bullet')

doc.add_page_break()

add_heading("6. System Screenshots & Result Analysis")

add_heading("6.1 Route Planner Interface", level=2)
add_paragraph("The following image demonstrates the system successfully calculating an optimal path between Delhi and Mumbai using Dijkstra's Algorithm.")
if os.path.exists('/Users/mac/.gemini/antigravity/brain/90e08bfa-22d0-430e-bc98-bfbf1932a50e/.user_uploaded/media_1786020561530.png'):
    doc.add_picture('/Users/mac/.gemini/antigravity/brain/90e08bfa-22d0-430e-bc98-bfbf1932a50e/.user_uploaded/media_1786020561530.png', width=Inches(6.0))

doc.add_page_break()

add_heading("6.2 Nearby & Emergency Locator", level=2)
add_paragraph("The image below showcases the localized spatial search. The Max-Heap spatial engine successfully located hospitals near Delhi.")
if os.path.exists('/Users/mac/.gemini/antigravity/brain/90e08bfa-22d0-430e-bc98-bfbf1932a50e/.user_uploaded/media_1786020583608.png'):
    doc.add_picture('/Users/mac/.gemini/antigravity/brain/90e08bfa-22d0-430e-bc98-bfbf1932a50e/.user_uploaded/media_1786020583608.png', width=Inches(6.0))

doc.add_page_break()

add_heading("6.3 Administrator Control Panel", level=2)
add_paragraph("This screenshot displays the sleek FlatLaf modernized UI applied to the administrative data injection panels.")
if os.path.exists('/Users/mac/.gemini/antigravity/brain/90e08bfa-22d0-430e-bc98-bfbf1932a50e/.user_uploaded/media_1786020612010.png'):
    doc.add_picture('/Users/mac/.gemini/antigravity/brain/90e08bfa-22d0-430e-bc98-bfbf1932a50e/.user_uploaded/media_1786020612010.png', width=Inches(6.0))

doc.add_page_break()

add_heading("7. Conclusion")
add_paragraph("The development of the Smart Route Planning System was a profound exercise in bridging theoretical computer science with practical software engineering. By successfully implementing and visualizing complex graph theories—ranging from A* heuristic pathfinding to Max-Heap spatial filtering—the project definitively proves the power of optimized Data Structures and Algorithms. The application not only meets all initial objectives but features a polished, deployable, and highly interactive user experience that parallels professional logistics software.")

doc.save('/Users/mac/Desktop/SmartRoutePlanner/Project_Report_Final_V3.docx')
